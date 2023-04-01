import docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication


doc = docx.Document()
full_names  = "Philip Mark Osodo"
email_address = "abc@xyz.com"
phone_number = "254704596323"
location = "Weiteithie"
occupation_title = "IT Personeel"
# linked_in = "philipmark.linkedin.com"
linked_in = ""
github = "mark.githib.com"
portifolio_link = ""
has_work_experience = True

uni_institution = "Maount Kenya University"
uni_degree = "Diploma"
uni_field = "Architecture"
uni_dates = "Jan 2016 to Dec 2019"

highschool_institution = "Baringo Boys High School"
highschool_dates = "Jan 2014 to Dec 2018"
highschool_grade = "C+"


primary_institution = "Makena School"
primary_dates = "Jan 2009 to Dec 2013"
primary_grade = "377"

skills_list = ["SQL","Python","(Pandas, NumPy & Matplotlib)"," Microsoft Excel","Power BI","Machine Learning using TensorFlow"]


# work_experiences = [("Job Title", "Location", "Date", "Duties 1\nDuty2\nDuty3")]
work_experiences = [("Quantitative Social Science Researcher", "American Institutes for Research", "March 2022 to Current",
                     "Develop conceptual frameworks that guide project work.\nManage and analyze large-scale data files for research."
                     "\nManage or lead projects and workstreams within projects\n"
                     "Design and support primary data collection\nBuilding data pipelines"),
                    ("Business System Analyst", "Tubali Limited", "January 2021 to March 2022",
                     "Data Collection\nData storage\nDatabase creation and maintenance\n"
                     "Project management and team collaboration\nBuilding data pipelines"),
                    ("IT Intern", "Ngamia Haulers Ltd", "June 2020 to December 2020",
                     "Computer Maintenance\nNetwork Maintenance\nIT support and troubleshooting\n"
                     "Project management and team collaboration\nData security and confidentiality protocols"),
                    ]

# referees = ["Full Names", "Role", "Institution", "phone/email"]

referees = [("Philip Oyier", "Lecturer", "Jomo kenyatta university of Agriculture and Technology", "+254714787478"),
            ("Margeret Mbulu", "IT Chairperson", "Ngamia Haulers Ltd", "+254774030108")]


sections = doc.sections
for section in sections:
    section.top_margin = docx.shared.Inches(0.5)
    section.bottom_margin = docx.shared.Inches(0.5)
    section.left_margin = docx.shared.Inches(1)
    section.right_margin = docx.shared.Inches(1)

font_styles = doc.styles
font_charstyle = font_styles.add_style('job_title_style', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(15)
font_object.bold = True


font_styles = doc.styles
font_charstyle = font_styles.add_style('Header_', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(17)

font_styles = doc.styles
font_charstyle = font_styles.add_style('section_headers', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(14)

font_styles = doc.styles
font_charstyle = font_styles.add_style('normal_text', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(12)


# INITIALIZATIONNNNN
heading = doc.add_heading('', level=1)
heading.add_run(f'{full_names.upper()}', style="Header_")
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER


job_title = doc.add_paragraph('')
job_title.add_run(f'{occupation_title.upper()}', style="job_title_style")
job_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# PERSONAL DETAILS

personal_details_header = doc.add_heading('', level=2)
personal_details_header.add_run("PERSONAL DETAILS", style = "section_headers")

phone = doc.add_paragraph("")
phone.add_run(f'Phone Number : +{phone_number}', style="normal_text")

email = doc.add_paragraph("")
email.add_run(f'Email Address: {email_address}', style="normal_text")

location_ = doc.add_paragraph("")
location_.add_run(f'Location: {location}', style="normal_text")

if linked_in != "":
    linkedin = doc.add_paragraph("")
    linkedin.add_run(f'Linked in: {linked_in}', style="normal_text")

if github != "":
    github_ = doc.add_paragraph("")
    github_.add_run(f'GitHub: {github}', style="normal_text")

if portifolio_link != "":
    portifolio = doc.add_paragraph("")
    portifolio.add_run(f'Portfolio: {portifolio_link}', style="normal_text")


overview = "A highly skilled data analyst with X years of experience in collecting, analyzing, and interpreting complex " \
           "data sets. Proficient in using statistical software such as R and Python, as well as data visualization " \
           "tools like Tableau and Power BI. Adept at communicating insights to both technical and non-technical " \
           "audiences, and skilled at collaborating with cross-functional teams to drive business outcomes."

overview_header = doc.add_heading('', level=2)
overview_header.add_run('OVERVIEW', style="section_headers")

the_overview = doc.add_paragraph("")
the_overview.add_run({overview}, style='normal_text')
the_overview.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


if has_work_experience is True:
    work_experience_header = doc.add_heading('', level=2)
    work_experience_header.add_run('WORK EXPERIENCE', style="section_headers")

    # [("Job Title", "Location", "Date", "Duties 1\nDuty2\nDuty3")]

    for job in work_experiences:
        job_title = job[0]
        job_location = job[1]
        job_dates = job[2]
        job_duties = job[3]

        job_duties = job_duties.split("\n")

        doc.add_paragraph(f"").add_run(f"{job_title.upper()} at {job_location.upper()}", style="normal_text").bold = True
        doc.add_paragraph(f"").add_run(f"{job_dates}", style="normal_text")
        doc.add_paragraph(f"").add_run(f"RESPONSIBILITIES", style="normal_text")

        bullet_format = doc.styles['List Bullet'].paragraph_format
        bullet_format.left_indent = docx.shared.Inches(0.5)
        bullet_format.space_after = docx.shared.Pt(16)
        for i in job_duties:
            bullet_item = doc.add_paragraph(f'{i}', style='List Bullet')

        doc.add_paragraph(f"").add_run(f"", style="normal_text")


doc.add_heading('', level=2).add_run('EDUCATION', style="section_headers")
doc.add_paragraph("").add_run({uni_institution.upper()}, style='normal_text').bold = True
doc.add_paragraph("").add_run(f'{uni_degree} in {uni_field}', style='normal_text')
doc.add_paragraph("").add_run(f'{uni_dates}', style='normal_text')

doc.add_paragraph("").add_run(f'', style='normal_text')
doc.add_paragraph("").add_run({highschool_institution.upper()}, style='normal_text').bold = True
doc.add_paragraph("").add_run(f' Kenya Certificate of Secondary Education(KCSE {highschool_grade})', style='normal_text')
doc.add_paragraph("").add_run(f'{highschool_dates}', style='normal_text')


doc.add_paragraph("").add_run(f'', style='normal_text')
doc.add_paragraph("").add_run({primary_institution.upper()}, style='normal_text').bold = True
if primary_grade != "":
    doc.add_paragraph("").add_run(f'Kenya Certificate of Primary Education (KCPE {primary_grade})', style='normal_text')
doc.add_paragraph("").add_run(f'{primary_dates}', style='normal_text')
doc.add_paragraph("").add_run(f'', style='normal_text')


doc.add_heading('', level=2).add_run('SKILLS', style="section_headers")
for i in skills_list:
    bullet_item = doc.add_paragraph(f'{i}', style='List Bullet')

doc.add_heading('', level=2).add_run('REFEREES', style="section_headers")
for i in referees:
    ref_name = i[0]
    ref_role = i[1]
    ref_company = i[2]
    ref_number = i[3]

    doc.add_paragraph("").add_run({ref_name.upper()}, style='normal_text').bold = True
    doc.add_paragraph("").add_run({ref_role.upper()}, style='normal_text')
    doc.add_paragraph("").add_run({ref_company.upper()}, style='normal_text')
    doc.add_paragraph("").add_run({ref_number.upper()}, style='normal_text')
    doc.add_paragraph("").add_run("", style='normal_text')


doc.save('resume.docx')


job_title = "Analyst"
company = 'Blueband'
poster_title = f"New applican for {job_title.lower()} position at {company}"

msg = MIMEMultipart()
msg['Subject'] = poster_title
msg['From'] = "noreply.jujaworks@gmail.com"
msg['To'] = 'noreply.jujaworks@gmail.com'

# add body text to the message
email_to_poster = "Testing 123"
body = MIMEText(email_to_poster)
msg.attach(body)

with open('resume.docx', 'rb') as f:
        # create the attachment object
        attachment = MIMEApplication(f.read(), _subtype='docx')
        attachment.add_header('Content-Disposition', 'attachment', filename='document.docx')
        # attach the file to the message
        msg.attach(attachment)

with smtplib.SMTP('smtp.gmail.com', 587) as smtp:
    smtp.starttls()
    password = "qsfqjbkroyolhyhi"

    smtp.login("noreply.jujaworks@gmail.com", password)
    smtp.send_message(msg)


